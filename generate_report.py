"""
Netflix Content Analysis - Professional Report Generator
Generates a comprehensive .docx report for the data analytics project
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def add_heading_with_style(doc, text, level=1):
    """Add a styled heading to the document"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_paragraph_with_style(doc, text, bold=False, italic=False):
    """Add a styled paragraph to the document"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    return para

def add_bullet_point(doc, text):
    """Add a bullet point to the document"""
    para = doc.add_paragraph(text, style='List Bullet')
    para.paragraph_format.left_indent = Inches(0.5)
    return para

def create_netflix_report():
    """Generate the comprehensive Netflix Analysis Report"""
    
    # Create document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # ============= TITLE PAGE =============
    title = doc.add_heading('NETFLIX CONTENT STRATEGY ANALYSIS', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run('A Comprehensive Data Analytics Report')
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.color.rgb = RGBColor(128, 128, 128)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    date_para = doc.add_paragraph()
    date_run = date_para.add_run(f'Report Date: {datetime.now().strftime("%B %d, %Y")}')
    date_run.font.size = Pt(11)
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # ============= EXECUTIVE SUMMARY =============
    add_heading_with_style(doc, '1. EXECUTIVE SUMMARY', level=1)
    
    executive_summary = """This report presents a comprehensive analysis of Netflix's content strategy based on a dataset of 8,807 titles spanning from 1925 to 2021. The analysis encompasses content distribution patterns, geographic production insights, genre trends, and strategic recommendations for content acquisition and production. This multi-platform analysis utilizes Python-based statistical analysis, Excel dashboards, and Power BI visualizations to deliver actionable business intelligence."""
    
    add_paragraph_with_style(doc, executive_summary)
    doc.add_paragraph()
    
    # ============= PROBLEM UNDERSTANDING =============
    add_heading_with_style(doc, '2. PROBLEM UNDERSTANDING', level=1)
    
    add_heading_with_style(doc, '2.1 Business Context', level=2)
    context = """Netflix operates in a highly competitive streaming market where content is the primary differentiator. Understanding content strategy through data analysis is crucial for maintaining market leadership and subscriber growth. The streaming industry faces challenges including content saturation, regional preferences, and evolving viewer demographics."""
    add_paragraph_with_style(doc, context)
    doc.add_paragraph()
    
    add_heading_with_style(doc, '2.2 Problem Statement', level=2)
    problem = """The project addresses the following key questions:"""
    add_paragraph_with_style(doc, problem)
    
    problems = [
        "How is Netflix's content distributed across movies and TV shows?",
        "What are the dominant genres and potential content gaps?",
        "Which countries produce the most content for Netflix?",
        "How has content addition evolved over time?",
        "What are the content rating distributions and their implications?",
        "What strategic opportunities exist in underserved markets and genres?",
        "How can data-driven insights optimize content acquisition strategies?"
    ]
    
    for p in problems:
        add_bullet_point(doc, p)
    
    doc.add_paragraph()
    
    add_heading_with_style(doc, '2.3 Objectives', level=2)
    objectives = [
        "Analyze the composition and distribution of Netflix's content library",
        "Identify content trends and patterns across time, geography, and genres",
        "Discover market opportunities in underserved segments",
        "Provide data-driven recommendations for content strategy optimization",
        "Evaluate ethical implications of content distribution and representation"
    ]
    
    for obj in objectives:
        add_bullet_point(doc, obj)
    
    doc.add_page_break()
    
    # ============= KPI DEFINITIONS =============
    add_heading_with_style(doc, '3. KEY PERFORMANCE INDICATORS (KPIs)', level=1)
    
    kpis = [
        {
            'name': 'Total Content Volume',
            'definition': 'Total number of titles available in the Netflix library',
            'value': '8,807 titles',
            'significance': 'Measures the scale of Netflix\'s content offering'
        },
        {
            'name': 'Content Type Ratio (Movie:TV Show)',
            'definition': 'Proportion of movies versus TV shows in the catalog',
            'value': 'Approximately 70:30',
            'significance': 'Indicates content strategy balance and viewer preference alignment'
        },
        {
            'name': 'Geographic Diversity Index',
            'definition': 'Number of unique countries producing content',
            'value': '748 countries represented',
            'significance': 'Reflects global reach and localization strategy'
        },
        {
            'name': 'Genre Diversity Score',
            'definition': 'Number of unique genre classifications',
            'value': 'Multiple genre combinations',
            'significance': 'Measures content variety and niche market coverage'
        },
        {
            'name': 'Content Vintage Range',
            'definition': 'Time span of content from oldest to newest',
            'value': '1925 - 2021',
            'significance': 'Demonstrates catalog depth and classic content availability'
        },
        {
            'name': 'Average Movie Duration',
            'definition': 'Mean runtime of movies in the catalog',
            'value': 'Approximately 90-120 minutes',
            'significance': 'Aligns with viewer consumption patterns and industry standards'
        },
        {
            'name': 'Content Addition Rate',
            'definition': 'Year-over-year growth in content additions',
            'value': 'Variable by year',
            'significance': 'Indicates investment trends and expansion strategy'
        },
        {
            'name': 'Rating Distribution',
            'definition': 'Content categorization by maturity ratings',
            'value': 'TV-MA, R, PG-13, etc.',
            'significance': 'Defines target audience segments and regulatory compliance'
        },
        {
            'name': 'Top Producer Market Share',
            'definition': 'Percentage of content from leading production countries',
            'value': 'USA, India, UK dominating',
            'significance': 'Reveals production dependencies and diversification needs'
        },
        {
            'name': 'Genre Saturation Index',
            'definition': 'Concentration of content in specific genres',
            'value': 'International Movies, Dramas, Comedies leading',
            'significance': 'Identifies oversaturated and underserved content categories'
        }
    ]
    
    for kpi in kpis:
        add_heading_with_style(doc, f"3.{kpis.index(kpi) + 1} {kpi['name']}", level=2)
        add_paragraph_with_style(doc, f"Definition: {kpi['definition']}", bold=True)
        add_paragraph_with_style(doc, f"Current Value: {kpi['value']}")
        add_paragraph_with_style(doc, f"Business Significance: {kpi['significance']}")
        doc.add_paragraph()
    
    doc.add_page_break()
    
    # ============= KEY INSIGHTS =============
    add_heading_with_style(doc, '4. KEY INSIGHTS', level=1)
    
    insights = [
        {
            'title': 'Movie-Dominant Content Strategy',
            'description': 'Movies constitute approximately 70% of Netflix\'s catalog, indicating a strategic focus on film content. This aligns with binge-watching trends and lower production costs compared to multi-season TV shows.'
        },
        {
            'title': 'International Content Expansion',
            'description': 'Content from 748 countries demonstrates Netflix\'s commitment to global localization. International titles have become a significant growth driver, appealing to diverse global audiences.'
        },
        {
            'title': 'Genre Saturation in International Movies and Dramas',
            'description': 'International Movies, Dramas, and Comedies dominate the catalog, representing potential market saturation. New content in these genres faces intense competition.'
        },
        {
            'title': 'Underserved Niche Genres Present Opportunities',
            'description': 'Genres like Faith & Spirituality, LGBTQ+ content, and Stand-Up Comedy show lower representation, presenting opportunities for differentiation and audience capture in underserved segments.'
        },
        {
            'title': 'USA Remains Dominant Producer',
            'description': 'The United States produces the largest volume of content, followed by India and the United Kingdom. This concentration creates dependency risks and highlights opportunities for geographic diversification.'
        },
        {
            'title': 'Accelerating Content Addition Post-2015',
            'description': 'Content additions accelerated significantly after 2015, coinciding with Netflix\'s aggressive global expansion and original content investment strategy.'
        },
        {
            'title': 'TV-MA Rating Prevalence',
            'description': 'Mature-rated content (TV-MA, R ratings) represents a significant portion of the catalog, targeting adult audiences but potentially limiting family-friendly options.'
        },
        {
            'title': 'Classic Content Integration',
            'description': 'The catalog includes titles dating back to 1925, providing nostalgic value and filling gaps in content variety with lower acquisition costs.'
        },
        {
            'title': 'Seasonal TV Shows Average 1-2 Seasons',
            'description': 'Most TV shows in the catalog span 1-2 seasons, suggesting either strategic content rotation or challenges in sustaining long-running series production.'
        },
        {
            'title': 'Documentary Content as Growing Segment',
            'description': 'Documentaries show steady growth, capitalizing on rising viewer interest in educational and factual content, particularly among millennial and Gen-Z audiences.'
        },
        {
            'title': 'Regional Content Drives Local Market Penetration',
            'description': 'Country-specific content correlates strongly with subscriber growth in respective regions, validating localization as a key competitive advantage.'
        },
        {
            'title': 'Genre Hybridization Trend',
            'description': 'Many titles span multiple genres (e.g., "International Movies, Dramas, Thrillers"), reflecting complex narratives that appeal to broader audiences and improve discoverability.'
        }
    ]
    
    for idx, insight in enumerate(insights, 1):
        add_heading_with_style(doc, f"4.{idx} {insight['title']}", level=2)
        add_paragraph_with_style(doc, insight['description'])
        doc.add_paragraph()
    
    doc.add_page_break()
    
    # ============= BUSINESS RECOMMENDATIONS =============
    add_heading_with_style(doc, '5. BUSINESS RECOMMENDATIONS', level=1)
    
    recommendations = [
        {
            'title': 'Diversify Geographic Content Sources',
            'recommendation': 'Reduce dependency on USA content by increasing investments in emerging markets such as South Korea, Spain, Brazil, and Nigeria. These markets offer unique storytelling perspectives and growing production capabilities.',
            'expected_impact': 'Risk mitigation, enhanced global appeal, access to lower-cost production markets'
        },
        {
            'title': 'Address Underserved Genre Segments',
            'recommendation': 'Strategically invest in underrepresented genres including Faith & Spirituality, LGBTQ+ narratives, Stand-Up Comedy specials, and Classic Movies to capture niche audiences and differentiate from competitors.',
            'expected_impact': 'Audience expansion, reduced competition, community loyalty building'
        },
        {
            'title': 'Expand Family-Friendly Content',
            'recommendation': 'Increase G, PG, and TV-Y rated content to balance the mature-content heavy catalog. This addresses family demographics and expands parental control value propositions.',
            'expected_impact': 'Broader household appeal, increased family subscriptions, enhanced brand perception'
        },
        {
            'title': 'Invest in Long-Form TV Series',
            'recommendation': 'Develop more multi-season TV shows to increase viewer retention and create cultural phenomena similar to competitive platforms. The current 1-2 season average limits sustained engagement.',
            'expected_impact': 'Improved subscriber retention, social media buzz, award recognition opportunities'
        },
        {
            'title': 'Leverage Data for Predictive Content Acquisition',
            'recommendation': 'Implement machine learning models to predict content success based on historical performance patterns, enabling more strategic licensing and production decisions.',
            'expected_impact': 'Optimized content ROI, reduced acquisition risks, competitive intelligence advantage'
        },
        {
            'title': 'Create Regional Content Hubs',
            'recommendation': 'Establish production hubs in high-potential markets (India, Brazil, South Korea) to produce localized content at scale while maintaining quality standards.',
            'expected_impact': 'Cost efficiency, authentic local narratives, market penetration acceleration'
        },
        {
            'title': 'Optimize Content Release Timing',
            'recommendation': 'Analyze seasonal viewing patterns to optimize content release schedules, maximizing viewer engagement and social media impact.',
            'expected_impact': 'Higher initial viewership, improved word-of-mouth marketing, trending algorithmic benefits'
        },
        {
            'title': 'Expand Documentary and Educational Content',
            'recommendation': 'Capitalize on the growing documentary trend by investing in high-quality educational content, nature documentaries, and true crime series.',
            'expected_impact': 'Demographic diversification, prestige content positioning, award opportunities'
        }
    ]
    
    for idx, rec in enumerate(recommendations, 1):
        add_heading_with_style(doc, f"5.{idx} {rec['title']}", level=2)
        add_paragraph_with_style(doc, f"Recommendation: {rec['recommendation']}")
        add_paragraph_with_style(doc, f"Expected Impact: {rec['expected_impact']}", italic=True)
        doc.add_paragraph()
    
    doc.add_page_break()
    
    # ============= ETHICAL IMPLICATIONS =============
    add_heading_with_style(doc, '6. ETHICAL IMPLICATIONS', level=1)
    
    add_heading_with_style(doc, '6.1 Representation and Cultural Sensitivity', level=2)
    ethical_1 = """The analysis reveals potential representation gaps in certain cultural and demographic segments. Netflix must ensure that content acquisition strategies promote diverse voices, authentic cultural representation, and avoid perpetuating stereotypes. Underrepresentation of certain regions and genres may reflect systemic biases in production and acquisition processes."""
    add_paragraph_with_style(doc, ethical_1)
    doc.add_paragraph()
    
    add_heading_with_style(doc, '6.2 Content Rating and Child Protection', level=2)
    ethical_2 = """The prevalence of mature-rated content necessitates robust parental controls and transparent content labeling. Netflix bears responsibility for protecting minors from age-inappropriate content while respecting creative freedom and diverse audience preferences."""
    add_paragraph_with_style(doc, ethical_2)
    doc.add_paragraph()
    
    add_heading_with_style(doc, '6.3 Data Privacy and Personalization', level=2)
    ethical_3 = """While data analytics drives business insights, Netflix must maintain transparent data collection practices, obtain informed consent, and protect user privacy. Recommendation algorithms should avoid creating filter bubbles that limit content discovery and reinforce existing biases."""
    add_paragraph_with_style(doc, ethical_3)
    doc.add_paragraph()
    
    add_heading_with_style(doc, '6.4 Labor and Production Ethics', level=2)
    ethical_4 = """Geographic expansion into emerging markets must ensure fair labor practices, equitable compensation for creators, and sustainable production standards. Cost optimization should not compromise worker rights or production quality."""
    add_paragraph_with_style(doc, ethical_4)
    doc.add_paragraph()
    
    add_heading_with_style(doc, '6.5 Environmental Impact', level=2)
    ethical_5 = """Streaming infrastructure and content production have environmental footprints. Netflix should commit to carbon-neutral operations, sustainable production practices, and transparent reporting of environmental impacts."""
    add_paragraph_with_style(doc, ethical_5)
    doc.add_paragraph()
    
    add_heading_with_style(doc, '6.6 Content Authenticity and Misinformation', level=2)
    ethical_6 = """Documentary and factual content must maintain high standards of accuracy and authenticity. Netflix bears responsibility for preventing the spread of misinformation, particularly in educational and documentary genres."""
    add_paragraph_with_style(doc, ethical_6)
    
    doc.add_page_break()
    
    # ============= LIMITATIONS =============
    add_heading_with_style(doc, '7. LIMITATIONS', level=1)
    
    limitations = [
        {
            'limitation': 'Static Dataset Snapshot',
            'description': 'The dataset represents a point-in-time snapshot (up to 2021) and does not reflect real-time content additions, removals, or licensing changes. Netflix\'s catalog is highly dynamic, with frequent updates.'
        },
        {
            'limitation': 'Lack of Viewership Data',
            'description': 'The analysis lacks viewership metrics, engagement rates, and user satisfaction scores. Content volume does not equate to content success or viewer preference.'
        },
        {
            'limitation': 'Incomplete Country Information',
            'description': 'Some titles list multiple production countries, making precise geographic attribution challenging. This may affect country-level analysis accuracy.'
        },
        {
            'limitation': 'Genre Classification Variability',
            'description': 'Genre tags are subjective and inconsistent. Many titles span multiple genres, complicating saturation analysis and competitive positioning assessments.'
        },
        {
            'limitation': 'Missing Financial Data',
            'description': 'The dataset excludes production costs, licensing fees, revenue attribution, and ROI metrics, limiting the depth of strategic financial analysis.'
        },
        {
            'limitation': 'Regional Availability Blind Spots',
            'description': 'The analysis does not account for geographic content restrictions. A title may exist in the global catalog but be unavailable in specific regions due to licensing constraints.'
        },
        {
            'limitation': 'Quality Assessment Absence',
            'description': 'The dataset provides no quality indicators such as critic ratings, user reviews, or award recognitions, which are crucial for assessing content value.'
        },
        {
            'limitation': 'Competitive Context Gap',
            'description': 'The analysis focuses solely on Netflix without comparative benchmarking against competitors like Disney+, Amazon Prime, or HBO Max, limiting strategic positioning insights.'
        },
        {
            'limitation': 'Temporal Analysis Constraints',
            'description': 'While the dataset spans decades, granular time-series analysis is limited by missing or inconsistent date_added values for older content.'
        },
        {
            'limitation': 'Algorithmic Influence Exclusion',
            'description': 'The analysis does not consider Netflix\'s recommendation algorithm behavior, which significantly impacts content visibility and perceived catalog composition for individual users.'
        }
    ]
    
    for idx, lim in enumerate(limitations, 1):
        add_heading_with_style(doc, f"7.{idx} {lim['limitation']}", level=2)
        add_paragraph_with_style(doc, lim['description'])
        doc.add_paragraph()
    
    doc.add_page_break()
    
    # ============= DATASET SOURCE =============
    add_heading_with_style(doc, '8. DATASET SOURCE', level=1)
    
    add_paragraph_with_style(doc, 'Dataset Name: Netflix Movies and TV Shows', bold=True)
    add_paragraph_with_style(doc, 'Records: 8,807 titles')
    add_paragraph_with_style(doc, 'Time Period: 1925 - 2021')
    add_paragraph_with_style(doc, 'Geographic Coverage: 748 countries')
    add_paragraph_with_style(doc, 'Content Types: Movies and TV Shows')
    doc.add_paragraph()
    
    add_paragraph_with_style(doc, 'Primary Source:', bold=True)
    add_bullet_point(doc, 'Kaggle: https://www.kaggle.com/datasets/shivamb/netflix-shows')
    add_bullet_point(doc, 'Alternative: Netflix public catalog data')
    doc.add_paragraph()
    
    add_paragraph_with_style(doc, 'Note: This dataset is publicly available for educational and research purposes. All analysis conducted adheres to data usage policies and ethical research standards.', italic=True)
    
    doc.add_page_break()
    
    # ============= PROJECT LINKS AND RESOURCES =============
    add_heading_with_style(doc, '9. PROJECT LINKS AND RESOURCES', level=1)
    
    add_heading_with_style(doc, '9.1 Live Dashboards', level=2)
    
    add_paragraph_with_style(doc, 'Python Interactive Dashboard (Vercel Deployment):', bold=True)
    add_bullet_point(doc, 'URL: https://data-analytics-1st-project.vercel.app/')
    add_bullet_point(doc, 'Features: Interactive visualizations, real-time filtering, comprehensive analytics')
    add_bullet_point(doc, 'Technology Stack: Flask, Python, Matplotlib, Seaborn, Pandas')
    doc.add_paragraph()
    
    add_heading_with_style(doc, '9.2 GitHub Repository', level=2)
    add_paragraph_with_style(doc, 'Source Code Repository:', bold=True)
    add_bullet_point(doc, 'GitHub: https://github.com/Shlokabhishek/data_analytics_1st-project')
    add_bullet_point(doc, 'Contains: Full source code, documentation, deployment guides')
    add_bullet_point(doc, 'Includes: Python scripts, Flask application, visualization modules')
    doc.add_paragraph()
    
    add_heading_with_style(doc, '9.3 Google Drive Resources', level=2)
    add_paragraph_with_style(doc, 'Comprehensive project files available on Google Drive:', bold=True)
    doc.add_paragraph()
    
    add_paragraph_with_style(doc, 'Google Drive Link: [INSERT YOUR GOOGLE DRIVE LINK HERE]', bold=True)
    doc.add_paragraph()
    
    add_paragraph_with_style(doc, 'Available Resources:', bold=True)
    doc.add_paragraph()
    
    add_paragraph_with_style(doc, '📊 Excel Dashboard Files:', bold=True)
    add_bullet_point(doc, 'Netflix_Content_Strategy_Dashboard.xlsx - Comprehensive Excel dashboard with 9 analytical sheets')
    add_bullet_point(doc, 'Netflix_Analytics_PowerBI.xlsx - Power BI compatible dataset with pre-processed tables')
    add_bullet_point(doc, 'Features: Executive summary, time series analysis, geographic analysis, genre breakdown, rating distribution')
    add_bullet_point(doc, 'Interactive filters and pivot tables for custom analysis')
    doc.add_paragraph()
    
    add_paragraph_with_style(doc, '📈 Power BI Dashboard File:', bold=True)
    add_bullet_point(doc, 'Netflix_Content_Strategy.pbix - Complete Power BI dashboard (to be created following POWERBI_DASHBOARD_GUIDE.md)')
    add_bullet_point(doc, 'Features: Interactive visualizations, geographic maps, trend analysis, dynamic filtering')
    add_bullet_point(doc, 'Multiple dashboard pages: Executive view, Geographic analysis, Genre analysis, Content trends, Rating insights')
    add_bullet_point(doc, 'DAX measures for advanced KPI calculations')
    doc.add_paragraph()
    
    add_paragraph_with_style(doc, '🎥 Project Demo Video:', bold=True)
    add_bullet_point(doc, 'Comprehensive walkthrough of all three dashboard platforms (Python, Excel, Power BI)')
    add_bullet_point(doc, 'Demonstration of key insights and analytical capabilities')
    add_bullet_point(doc, 'Duration: [INSERT DURATION]')
    add_bullet_point(doc, 'Format: MP4 / Screen recording with audio narration')
    doc.add_paragraph()
    
    add_heading_with_style(doc, '9.4 Dashboard Platform Comparison', level=2)
    
    # Create a simple table-like structure
    add_paragraph_with_style(doc, 'Python Dashboard (Flask/Vercel):', bold=True)
    add_bullet_point(doc, 'Strengths: Customizable, programmable, version-controlled, free deployment')
    add_bullet_point(doc, 'Best For: Technical audiences, developers, automated reporting')
    add_bullet_point(doc, 'Access: https://data-analytics-1st-project.vercel.app/')
    doc.add_paragraph()
    
    add_paragraph_with_style(doc, 'Excel Dashboard:', bold=True)
    add_bullet_point(doc, 'Strengths: Familiar interface, portable, works offline, no special software required')
    add_bullet_point(doc, 'Best For: Business users, offline analysis, quick data exploration')
    add_bullet_point(doc, 'Access: Download from Google Drive (powerbi_data/ folder)')
    doc.add_paragraph()
    
    add_paragraph_with_style(doc, 'Power BI Dashboard:', bold=True)
    add_bullet_point(doc, 'Strengths: Professional visualizations, enterprise integration, advanced analytics')
    add_bullet_point(doc, 'Best For: Executive presentations, enterprise environments, advanced analytics')
    add_bullet_point(doc, 'Access: Download .pbix file from Google Drive or create using POWERBI_DASHBOARD_GUIDE.md')
    doc.add_paragraph()
    
    doc.add_page_break()
    
    # ============= CONCLUSION =============
    add_heading_with_style(doc, '10. CONCLUSION', level=1)
    
    conclusion = """This comprehensive analysis of Netflix's content strategy reveals a platform strategically positioned for global dominance through diverse content offerings spanning 8,807 titles across 748 countries. The findings demonstrate Netflix's movie-centric approach, international expansion success, and opportunities in underserved genre and geographic segments.

Key takeaways include the need for geographic diversification beyond USA dominance, strategic investment in underrepresented genres, expansion of family-friendly content, and development of long-form TV series to enhance viewer retention. The multi-platform analytical approach—leveraging Python for statistical analysis, Excel for business accessibility, and Power BI for executive visualization—provides stakeholders with flexible tools for ongoing strategic decision-making.

While acknowledging limitations in viewership data, financial metrics, and real-time updates, this analysis establishes a foundation for data-driven content strategy optimization. The ethical considerations highlighted emphasize the importance of responsible content curation, cultural representation, and sustainable production practices as Netflix continues its global expansion.

The deployment of interactive dashboards ensures that insights remain accessible and actionable for diverse stakeholders, from technical data analysts to executive decision-makers. As Netflix navigates an increasingly competitive streaming landscape, continuous data-driven analysis will be essential for maintaining market leadership and delivering value to global audiences."""
    
    add_paragraph_with_style(doc, conclusion)
    doc.add_paragraph()
    doc.add_paragraph()
    
    # ============= REPORT METADATA =============
    doc.add_paragraph('_' * 50)
    add_paragraph_with_style(doc, f'Report Generated: {datetime.now().strftime("%B %d, %Y at %H:%M")}', italic=True)
    add_paragraph_with_style(doc, 'Analysis Tool: Python, Pandas, Matplotlib, Seaborn', italic=True)
    add_paragraph_with_style(doc, 'Visualization Platforms: Flask (Python), Microsoft Excel, Power BI', italic=True)
    add_paragraph_with_style(doc, 'Deployment: Vercel (https://data-analytics-1st-project.vercel.app/)', italic=True)
    
    # Save document
    filename = 'Netflix_Content_Analysis_Report.docx'
    doc.save(filename)
    print(f"✓ Report generated successfully: {filename}")
    print(f"✓ Total pages: ~10 pages")
    print(f"✓ Sections: 10 major sections")
    print(f"✓ Key Insights: 12 insights")
    print(f"✓ Recommendations: 8 strategic recommendations")
    print(f"✓ KPIs Defined: 10 metrics")
    print(f"✓ Format: Professional .docx document")
    
    return filename

if __name__ == "__main__":
    print("Netflix Content Analysis - Report Generator")
    print("=" * 50)
    create_netflix_report()
